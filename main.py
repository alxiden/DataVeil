import time
import tkinter as tk
import os
import shutil
from tkinter import messagebox
from docx import Document
import openpyxl
import csv
import PyPDF2
from PyPDF2 import PdfReader, PdfWriter
import extract_msg
import win32com.client
import fitz
import re
from html.parser import HTMLParser


class DataVeil:
    def __init__(self, root):
        self.string_storage = []

        self.root = root
        self.root.title("DataVeil")

        self.title_label = tk.Label(root, text="DataVeil")
        self.title_label.grid(row=0, column=0, columnspan=2)

        self.folder_label = tk.Label(root, text="Folder Location:")
        self.folder_label.grid(row=1, column=0)

        self.folder_entry = tk.Entry(root, width=50)
        self.folder_entry.grid(row=1, column=1)

        self.strings_label = tk.Label(root, text="Strings to Redact:")
        self.strings_label.grid(row=2, column=0)

        self.strings_entry = tk.Entry(root, width=50)
        self.strings_entry.grid(row=2, column=1)

        self.redact_emails_var = tk.BooleanVar()
        self.redact_money_var = tk.BooleanVar()
        self.redact_links_var = tk.BooleanVar()

        self.redact_emails_checkbox = tk.Checkbutton(root, text="Redact emails", variable=self.redact_emails_var)
        self.redact_emails_checkbox.grid(row=3, column=0, sticky='w')

        self.redact_money_checkbox = tk.Checkbutton(root, text="Redact Money", variable=self.redact_money_var)
        self.redact_money_checkbox.grid(row=3, column=1, sticky='w')

        self.redact_links_checkbox = tk.Checkbutton(root, text="Redact Links", variable=self.redact_links_var)
        self.redact_links_checkbox.grid(row=4, column=0, sticky='w')

        self.redact_button = tk.Button(root, text="Redact", command=self.files)
        self.redact_button.grid(row=5, column=0, columnspan=2)

    def files(self):
        folder = self.folder_entry.get()
        redacted_folder = os.path.join(folder, "Redacted")
        try:
            self.text_var()
            if not os.path.exists(redacted_folder):
                os.makedirs(redacted_folder)
            files = os.listdir(folder)
            for file in files:
                file_path = os.path.join(folder, file)
                if os.path.isfile(file_path):
                    shutil.copy(file_path, redacted_folder)
            redacted_files = os.listdir(redacted_folder)
            self.fileTypes(redacted_files, redacted_folder)
        except FileNotFoundError:
            self.show_error_popup(f"Folder '{folder}' not found.")
        except Exception as e:
            self.show_error_popup(f"An error occurred: {e}")
        
    def text_var(self):
        self.string_storage = []
        self.strings = self.strings_entry.get().split(',')
        for string in self.strings:
            self.string_storage.append(string.lower())
            self.string_storage.append(string.upper())
            self.string_storage.append(string.capitalize())
            self.string_storage.append(string.title())
            self.string_storage.append(string)
        #print(self.string_storage)

        
    def show_popup(self, message):
        messagebox.showinfo("Redaction Complete", message)
    
    def show_error_popup(self, message):
        messagebox.showerror("Error", message)

    def fileTypes(self, files, folder):
        for file in files:
            print(f"Processing file: {file}")
            file_path = os.path.join(folder, file)
            #print(file_path)
            if file.endswith('.txt'):
                self.redact_txt(file_path)
            elif file.endswith('.csv'):
                self.redact_csv(file_path)
            elif file.endswith('.xlsx'):
                self.redact_xlsx(file_path)
            elif file.endswith('.docx'):
                self.redact_docx(file_path)
            elif file.endswith('.pdf'):
                self.redact_pdf(file_path)
            elif file.endswith('.msg'):
                self.redact_msg(file_path)
            elif file.endswith('.HTML') or file.endswith('.html'):
                self.redact_txt(file_path)
            else:
               self.show_error_popup(f"File type {file} not supported")
        self.show_popup("Redaction complete for all files")

    def redact_txt(self, file):
        #print(file)
        strings_to_redact = self.string_storage
        try:
            with open(file, 'r') as f:
                content = f.read()
            for string in strings_to_redact:
                content = content.replace(string, "Redacted")
            with open(file, 'w') as f:
                f.write(content)
        except Exception as e:
            self.show_error_popup(f"An error occurred while processing the file '{file}': {e}")

    def redact_csv(self, file):
        #print(file)
        strings_to_redact = self.string_storage
        try:
            with open(file, 'r', newline='') as f:
                reader = csv.reader(f)
                rows = list(reader)

            for i, row in enumerate(rows):
                for j, cell in enumerate(row):
                    for string in strings_to_redact:
                        if string in cell:
                            rows[i][j] = cell.replace(string, "Redacted")

            with open(file, 'w', newline='') as f:
                writer = csv.writer(f)
                writer.writerows(rows)
        except Exception as e:
            self.show_error_popup(f"An error occurred while processing the file '{file}': {e}")
    
    def redact_xlsx(self, file):
        #print(file)
        strings_to_redact = self.string_storage
        try:
            wb = openpyxl.load_workbook(file)
            for sheet in wb.worksheets:
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value is not None:
                            cell_value_str = str(cell.value)
                            for string in strings_to_redact:
                                if string in cell_value_str:
                                    cell_value_str = cell_value_str.replace(string, "Redacted")
                            cell.value = cell_value_str
            wb.save(file)
        except Exception as e:
            self.show_error_popup(f"An error occurred while processing the file '{file}': {e}")
    
    def redact_docx(self, file):
        # print(file)
        strings_to_redact = self.string_storage
        redact_emails = self.redact_emails_var.get()
        redact_money = self.redact_money_var.get()
        email_pattern = r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+"
        money_pattern = r"£\d+"
        money_pattern_2 = r"£\d+\.\d{2}"
        money_pattern_3 = r"\b\d+[Kk]"
        link_pattern = r"https?://[^\s]+"
        probate_pattern = r"\d+\.\d{3}"
        exception_email = "vicki.wheelhouse@kctrust.co.uk"
        try:
            doc = Document(file)
            for paragraph in doc.paragraphs:
                text = paragraph.text
                # Redact custom strings
                for string in strings_to_redact:
                    if string in text:
                        text = text.replace(string, "Redacted")
                if redact_emails:
                    # Find all emails
                    found_emails = re.findall(email_pattern, text)
                    for email in found_emails:
                        if email.lower() != exception_email:
                            text = text.replace(email, "Redacted")
                if redact_money:
                    text = re.sub(money_pattern, "Redacted", text)
                    text = re.sub(money_pattern_2, "Redacted", text)
                    text = re.sub(money_pattern_3, "Redacted", text)
                if self.redact_links_var.get():
                    text = re.sub(link_pattern, "Redacted", text)
                if re.search(probate_pattern, text):
                    text = re.sub(probate_pattern, "Redacted", text)
                paragraph.text = text
            doc.save(file)
        except ValueError as ve:
            self.show_error_popup(ve)
        except Exception as e:
            self.show_error_popup(f"An error occurred while processing the file '{file}': {e}")

    def redact_pdf(self, file):
        strings_to_redact = self.string_storage
        try:
            doc = fitz.open(file)
            for page in doc:
                for string in strings_to_redact:
                    text_instances = page.search_for(string)
                    for inst in text_instances:
                        page.add_redact_annot(inst, fill=(0, 0, 0))  # Redact with black color
                    page.apply_redactions()
            temp_file = file.replace('.pdf', '_redacted.pdf')
            doc.save(temp_file, garbage=4, deflate=True)  # Apply redactions and save to a new file
            doc.close()
            os.replace(temp_file, file)  # Replace the original file with the redacted file
        except Exception as e:
            self.show_error_popup(f"An error occurred while processing the file '{file}': {e}")
    
    def convert_msg_to_docx(self, msg_file, redacted_folder):
        
        class HTMLTextExtractor(HTMLParser):
            def __init__(self):
                super().__init__()
                self.result = []
            def handle_data(self, d):
                self.result.append(d)
            def get_text(self):
                return ''.join(self.result)
        def sanitize(s):
            # Remove control characters and ensure XML compatibility
            return ''.join(c for c in str(s) if c.isprintable() and ord(c) not in range(0,32))

        msg = extract_msg.Message(msg_file)
        # print(dir(msg))  # Debugging line to check available attributes
        # Format date for filename (remove invalid chars)
        date_str = sanitize(msg.date).replace(':', '-').replace('/', '-').replace(' ', '_')
        sender_str = re.sub(r'[^a-zA-Z0-9]', '_', sanitize(msg.sender))
        docx_filename = f"{date_str}{sender_str}.docx"
        docx_file = os.path.join(redacted_folder, docx_filename)
        doc = Document()
        doc.add_heading('Email Information', level=1)
        doc.add_paragraph(f"Subject: {sanitize(msg.subject)}")
        doc.add_paragraph(f"From: {sanitize(msg.sender)}")
        doc.add_paragraph(f"To: {sanitize(msg.to)}")
        doc.add_paragraph(f"Date: {sanitize(msg.date)}")
        doc.add_heading('Body', level=1)
        def ensure_str(val):
            if isinstance(val, bytes):
                try:
                    return val.decode('utf-8')
                except Exception:
                    return val.decode('latin1', errors='ignore')
            return str(val)

        body_content = sanitize(ensure_str(msg.body))
        if body_content == "None":
            # Try to get the full message body or html body if available
            if hasattr(msg, 'messageBody') and msg.messageBody:
                body_content = sanitize(ensure_str(msg.messageBody))
            elif hasattr(msg, 'htmlBody') and msg.htmlBody:
                # Remove <style> and <head> sections from HTML
                html = ensure_str(msg.htmlBody)
                html = re.sub(r'<style.*?>.*?</style>', '', html, flags=re.DOTALL|re.IGNORECASE)
                html = re.sub(r'<head.*?>.*?</head>', '', html, flags=re.DOTALL|re.IGNORECASE)
                parser = HTMLTextExtractor()
                parser.feed(html)
                text = parser.get_text()
            else:
                body_content = "No body content found."
        doc.add_paragraph(body_content)
        doc.save(docx_file)
        msg.close()
        return docx_file
    
    def redact_msg(self, file):
        # Save redacted emails in the Redacted folder
        redacted_folder = os.path.dirname(file)
        try:
            docx_file = self.convert_msg_to_docx(file, redacted_folder)
            self.redact_docx(docx_file)
            time.sleep(2)
            os.remove(file)
        except Exception as e:
            self.show_error_popup(f"An error occurred while processing the file '{file}': {e}")


def main():
    root = tk.Tk()
    app = DataVeil(root)
    root.mainloop()

if __name__ == '__main__':
    main()